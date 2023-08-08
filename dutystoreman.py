import random
from datetime import date, timedelta
from docx import Document

# List of storeman names
storeman_names = ["Ruben", "Danial", "Khai", "Jude", "Afiq"]


# Function to generate the duty roster for the whole month with restrictions
def generate_duty_roster_with_restrictions(year, month, unavailable_dates):
    start_date = date(year, month, 1)
    end_date = start_date.replace(month=month + 1) - timedelta(days=1)
    duty_roster = {}
    weekly_assignments = {}

    current_date = start_date
    while current_date <= end_date:
        if current_date.weekday() < 5:  # Weekday (Monday to Friday)
            week_number = current_date.isocalendar()[1]
            available_storemen = [name for name in storeman_names if
                                  name not in unavailable_dates.get(current_date, [])]

            if week_number not in weekly_assignments:
                weekly_assignments[week_number] = random.sample(available_storemen, len(available_storemen))

            storeman = weekly_assignments[week_number].pop(0)
            duty_roster[current_date.strftime("%d-%m-%Y")] = storeman
        else:  # Weekend (Saturday and Sunday)
            duty_roster[current_date.strftime("%d-%m-%Y")] = "Weekend"

        current_date += timedelta(days=1)

    return duty_roster


# Generate the duty roster for August 2023 with restrictions
unavailable_dates = {
    #date(2023, 8, 5): ["Ruben"],  # Ruben is unavailable on 05-08-2023
}

duty_roster = generate_duty_roster_with_restrictions(2023, 9, unavailable_dates)

# Create a Word document
doc = Document()
doc.add_heading('Duty Roster - September 2023', level=1)

# Add a table to the document
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Add column headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Date'
hdr_cells[1].text = 'Storeman'

# Add duty roster data to the table
for date, storeman in duty_roster.items():
    row_cells = table.add_row().cells
    row_cells[0].text = date
    row_cells[1].text = storeman

# Save the Word document
doc.save("duty_roster.docx")
